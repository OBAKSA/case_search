# -*- coding:utf-8 -*-
import os
from collections import namedtuple
from urllib.request import urlopen
from urllib import parse
from bs4 import BeautifulSoup as bs
from docx import Document

__BASE_URL = "www.law.go.kr"
Pan = namedtuple("Pan", "pansi yozi jomun refpan allcon num")


def __get_bs_obj(url):
    '''
    url (str)
    return beautifulSoup object
    '''
    url = "http://"+url
    html = urlopen(url)
    return bs(html, 'html.parser', from_encoding='utf-8')


def __parse_pan(txt):
    txt_lines = txt.splitlines()
    next_lines = txt_lines[1:]
    contents = [None for i in range(0, 5)]

    idx = 0
    for line, next_line in zip(txt_lines, next_lines):
        if '【판시사항】' in line:
            contents[0] = next_line
        elif '【판결요지】' in line:
            contents[1] = next_line
        elif '【참조조문】' in line:
            contents[2] = next_line
        elif '【참조판례】' in line:
            contents[3] = next_line
        elif '【전문】' in line:
            allcon = next_line
            for con_line in txt_lines[idx+2:]:
                allcon += con_line+"\n"
            contents[4] = allcon
            break
        idx += 1
    return contents


def __split_number_list(txt):
    otxt = ""
    for c, next_c in zip(txt, txt[1:]):
        if (c == "[" or c == "(") and next_c.isnumeric():
            otxt += "\n"
        elif c in "가나라마바사아자차카타파하" and next_c == ".":
            otxt += "\n"
        otxt += c
    otxt += txt[-1]
    return otxt


def get_pan(pan_num):
    url = "{}/판례/({})".format(__BASE_URL, pan_num)
    website = __get_bs_obj(parse.quote(url))  # handle the hangeul url
    pansrc = website.body.find('iframe', {'name': 'lawService'})
    if pansrc is None:
        return Pan(None, None, None, None, None, pan_num)
    pansrc_url = pansrc['src']
    website = __get_bs_obj(__BASE_URL+"/"+pansrc_url)
    contents = website.body.find('div', {'class': 'pgroup'}).text
    contents = __parse_pan(contents)
    contents = [__split_number_list(txt) for txt in contents]
    contents.append(pan_num)
    return Pan(*contents)


def save_pans(pans, ofilename, p_check=True, y_check=True, j_check=True, r_check=True, a_check=True):
    '''
    pans (list of Pan objects)
    ofilename (str) : file path to save pans as docx file
    '''
    doc = Document()
    pan_nums = ""
    for pan in pans:
        pan_nums += pan.num+"\n"

    doc.add_heading("판례번호", level=1)
    doc.add_paragraph(pan_nums)
    doc.add_page_break()

    for pan in pans:
        doc.add_heading(pan.num, level=1)

        if pan.pansi or pan.yozi or pan.jomun or pan.refpan or pan.allcon:
            if p_check and pan.pansi:
                doc.add_heading("판시사항", level=2)
                doc.add_paragraph(pan.pansi)
            if y_check and pan.yozi:
                doc.add_heading("판결요지", level=2)
                doc.add_paragraph(pan.yozi)
            if j_check and pan.jomun:
                doc.add_heading("참조조문", level=2)
                doc.add_paragraph(pan.jomun)
            if r_check and pan.refpan:
                doc.add_heading("참조판례", level=2)
                doc.add_paragraph(pan.refpan)
            if a_check and pan.allcon:
                doc.add_heading("전문", level=2)
                doc.add_paragraph(pan.allcon)
        else:
            doc.add_heading("내용이 없습니다", level=2)

        doc.add_page_break()

    doc.save(ofilename)
    print("save: {}".format(os.path.abspath(ofilename)))


if __name__ == "__main__":
    get_pan("2012다13507")
